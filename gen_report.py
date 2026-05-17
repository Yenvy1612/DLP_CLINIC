import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text):
    p = doc.add_paragraph(text)
    return p

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading("BÁO CÁO DỰ ÁN", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("HỆ THỐNG QUẢN LÝ PHÒNG KHÁM A-CARE TÍCH HỢP DLP", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # CHƯƠNG 1
    add_heading(doc, "CHƯƠNG 1: MỞ ĐẦU", 1)

    add_heading(doc, "1.1. Bối cảnh đề tài", 2)
    add_para(doc, "Trong bối cảnh chuyển đổi số y tế diễn ra mạnh mẽ, thông tin sức khỏe (hồ sơ bệnh án, căn cước công dân, số điện thoại) ngày càng được lưu trữ số hóa nhiều hơn. Đi kèm với sự tiện lợi là rủi ro về việc rò rỉ dữ liệu cá nhân, thất thoát thông tin nhạy cảm từ các tác nhân bên trong (nhân viên y tế) và bên ngoài. Vì vậy, việc áp dụng công nghệ chống thất thoát dữ liệu (DLP - Data Loss Prevention) vào các hệ thống phòng khám là vô cùng cấp thiết.")

    add_heading(doc, "1.2. Lý do chọn đề tài", 2)
    add_para(doc, "Đề tài được lựa chọn dựa trên nhu cầu thực tiễn về việc bảo mật thông tin y tế theo các chuẩn mực bảo mật thông tin toàn cầu (như HIPAA). Hệ thống A-Care không chỉ là một giải pháp quản lý lịch hẹn thông thường mà còn là một cơ chế phòng vệ an toàn dữ liệu nhiều lớp (từ máy chủ đến thiết bị di động), giải quyết bài toán kiểm soát hành vi người dùng (đặc biệt là bác sĩ) khi xuất hay tải các thông tin bệnh án.")

    add_heading(doc, "1.3. Mục tiêu của đề tài", 2)
    add_para(doc, "1. Xây dựng nền tảng Quản lý phòng khám (A-Care) hoàn chỉnh với các chức năng đặt lịch, xem hồ sơ bệnh án.\n"
                  "2. Nghiên cứu và áp dụng cơ chế DLP phân tán (Mobile Agent và Backend API).\n"
                  "3. Tích hợp phân tích hành vi HTTP (HTTP Behavior Check) và ghi nhận nhật ký thất thoát (DLP Logs) nhằm giảm thiểu rủi ro bảo mật.")

    add_heading(doc, "1.4. Đối tượng và phạm vi nghiên cứu", 2)
    add_para(doc, "- Đối tượng: Quy trình khám chữa bệnh tại phòng khám đa khoa, dữ liệu cá nhân và thông tin y tế nhạy cảm.\n"
                  "- Phạm vi: Hệ thống Backend (Spring Boot) và Mobile App (Android Kotlin). Giới hạn ở việc kiểm soát hành vi truy cập API, đánh giá rủi ro theo quy tắc tĩnh, và Agent che lấp dữ liệu khi xuất PDF tại ứng dụng di động.")

    add_heading(doc, "1.5. Phương pháp nghiên cứu", 2)
    add_para(doc, "Nghiên cứu theo hướng phát triển phần mềm kết hợp thực nghiệm an toàn thông tin. Áp dụng kỹ thuật AOP (Aspect-Oriented Programming) trên Backend để đánh giá HTTP Behavior. Thiết kế Mobile Agent chạy ngầm trên Android để kiểm soát quyền và che lấp dữ liệu (masking).")

    add_heading(doc, "1.6. Cấu trúc của báo cáo", 2)
    add_para(doc, "Báo cáo bao gồm 5 chương từ mở đầu, cơ sở lý thuyết, phân tích thiết kế, cài đặt thử nghiệm cho đến tổng kết và hướng phát triển.")

    # CHƯƠNG 2
    add_heading(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    add_para(doc, "- Data Loss Prevention (DLP): Là giải pháp giúp phát hiện và ngăn chặn việc rò rỉ dữ liệu nhạy cảm ra ngoài.\n"
                  "- Phân tích hành vi (User Behavior Analytics): Xác định rủi ro qua tần suất, thời gian và dung lượng truy cập.\n"
                  "- Công nghệ Backend: Java Spring Boot, MySQL, Spring Security, JWT, AspectJ (AOP), Redis (quản lý rate-limit).\n"
                  "- Công nghệ Mobile: Kotlin, Android Architecture Components (MVVM, Navigation), Room Database (Agent storage).")

    # CHƯƠNG 3
    add_heading(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_para(doc, "Hệ thống A-Care thiết kế theo mô hình client-server, bao gồm Clinic Mobile App kết nối với Spring Boot Backend thông qua REST API. Đặc biệt, Mobile App tích hợp sẵn một Mobile Agent đóng vai trò phân tích bảo mật ngay trên thiết bị người dùng.")

    add_heading(doc, "3.1. Tổng quan Kiến trúc và Mobile Agent", 2)
    add_para(doc, "Clinic Mobile App không chỉ phục vụ đặt lịch, xem hồ sơ, mà còn bao gồm Module Agent giúp:\n"
                  "- Đăng ký thiết bị Android với backend.\n"
                  "- Gửi heartbeat/status định kỳ.\n"
                  "- Đồng bộ policy DLP từ backend.\n"
                  "- Quét dữ liệu nhạy cảm trong các hành vi nhập form, copy, và export.\n"
                  "- Đẩy log sự kiện về backend hoặc lưu trữ offline bằng Room Database.")

    add_heading(doc, "3.2. Mô tả các luồng nghiệp vụ cốt lõi", 2)

    add_heading(doc, "3.2.1. Luồng khởi động và Splash", 3)
    add_para(doc, "Khi người dùng mở app, phương thức `AgentInitializer.init()` được gọi trong class `App.kt` để khởi tạo toàn bộ các module của Agent (RetrofitClient, AgentDatabase, DlpScanner, EventTracker, v.v.).\n"
                  "Sau đó, SplashActivity kiểm tra token trong SessionManager:\n"
                  "- Nếu tồn tại token hợp lệ: Chuyển thẳng vào MainActivity.\n"
                  "- Nếu không tồn tại: Chuyển sang LoginActivity.")

    add_heading(doc, "3.2.2. Luồng đăng nhập và đăng ký", 3)
    add_para(doc, "Người dùng nhập Username/Password tại LoginActivity -> Gọi API đăng nhập.\n"
                  "Backend (AuthController) trả về JWT Token. Sau khi lưu Token vào SessionManager, Mobile Agent sẽ lập tức thực hiện 3 tác vụ bảo mật:\n"
                  "1. Register Device: Gửi thông tin thiết bị (deviceId, platform, OS) lên backend.\n"
                  "2. Sync Policy: Lấy các luật regex và keyword nhạy cảm từ backend.\n"
                  "3. Send Heartbeat: Báo cáo trạng thái hoạt động đầu tiên.\n"
                  "Sau khi hoàn tất, giao diện chuyển sang MainActivity.")

    add_heading(doc, "3.2.3. Luồng DLP Scan trong Form và Copy/Export", 3)
    add_para(doc, "Đây là luồng quan trọng nhất để bảo vệ rò rỉ dữ liệu phía Client:\n"
                  "- Luồng Copy/Export: Khi người dùng bấm Copy hoặc Export PDF, dữ liệu trước tiên đi qua `DlpScanner.scan()`. Nếu phát hiện CCCD, số điện thoại, hoặc từ khóa cấm (HIV, Tuyệt mật, Ung thư...), app sẽ chặn hành động (Block) và tạo event gửi về Backend (ví dụ: `COPY_PATIENT_DATA` với `severity=HIGH`).\n"
                  "- Luồng Masking: Một số hành động cho phép xuất PDF nhưng dữ liệu nhạy cảm sẽ bị che lấp (Masking) bởi hàm `MaskingUtil.mask()`, biến các chuỗi số CCCD thành định dạng `***`.")

    add_heading(doc, "3.2.4. Luồng xử lý Queue Offline (Mất mạng)", 3)
    add_para(doc, "Hệ thống DLP hoạt động bất chấp trạng thái mạng của thiết bị di động:\n"
                  "Khi phát sinh vi phạm, `AgentEventTracker` đưa event vào `EventQueueRepository`. Dữ liệu được lưu xuống Room Database cục bộ (`agent_events`). `EventSyncWorker` chạy ngầm, nếu phát hiện có mạng sẽ đẩy toàn bộ log lên API `/api/agent-events`. Nếu mất mạng, event tiếp tục được lưu giữ cho đến lần đồng bộ tiếp theo.")

    add_heading(doc, "3.3. Module Behavior Check và DLP Logs (Phía Backend)", 2)
    add_para(doc, "Ở phía Backend, mọi request truy xuất bệnh án được intercept bởi `BehaviorAspect` và `BehaviorTrackingFilter`:\n"
                  "- Rule Engine: Kiểm tra nếu truy cập ngoài giờ hành chính, vượt quá Rate limit hoặc số lượng tải. Nếu vi phạm, sẽ sinh ra SecurityException chặn request.\n"
                  "- DLP Logs: Các sự kiện từ Mobile Agent hoặc từ Filter Backend sẽ được `AgentEventService` đánh giá và phân loại vào `dlp_logs` với các mức độ LOW, MEDIUM, HIGH, CRITICAL. Admin Dashboard có thể dựa vào dữ liệu này để thống kê bất thường (Anomaly Detection).")

    # CHƯƠNG 4
    add_heading(doc, "CHƯƠNG 4: CÀI ĐẶT VÀ THỬ NGHIỆM", 1)
    add_para(doc, "Hệ thống đã được thiết lập thành công trên môi trường localhost. Các package của Android (ui, agent, data) và Spring Boot (Controller, Service, Filter) hoạt động trơn tru.")
    
    add_heading(doc, "4.1. Kịch bản thử nghiệm end-to-end", 2)
    add_para(doc, "Bước 1: Mở app, đăng nhập bằng tài khoản Bác sĩ.\n"
                  "Bước 2: App tự động đăng ký thiết bị với Backend. Admin kiểm tra thấy thiết bị ANDROID đã trực tuyến.\n"
                  "Bước 3: Bác sĩ tải danh sách bệnh án. App gửi event `VIEW_MEDICAL_RECORD_LIST`.\n"
                  "Bước 4: Bác sĩ thử bấm Export PDF có chứa CCCD. `MaskingUtil` nhận diện thành công CCCD/Email/SĐT và thay bằng `***` trên file xuất ra.\n"
                  "Bước 5: Bác sĩ bấm Copy đoạn văn bản chứa CCCD hoặc từ khóa nhạy cảm. Ứng dụng từ chối Copy, đồng thời gửi ngay event `COPY_PATIENT_DATA` kèm cảnh báo HIGH về Server.\n"
                  "Bước 6: Backend nhận log, phân tích rủi ro và đẩy lên bảng theo dõi bảo mật của Admin.")

    # CHƯƠNG 5
    add_heading(doc, "CHƯƠNG 5: TỔNG KẾT VÀ ĐÁNH GIÁ", 1)
    add_para(doc, "Hệ thống DLP Clinic đã đáp ứng tốt yêu cầu thực tế, tạo ra vòng bảo vệ kép từ API Backend đến Mobile Agent. Việc phân lớp quản lý log và quy tắc bảo mật (Policy Syncing) mang lại sự linh hoạt cho nhà quản trị mạng, đảm bảo không có bất kì dữ liệu nhạy cảm nào bị đưa ra ngoài một cách vô ý hoặc cố tình mà không có vết tích (Audit trail).\n"
                  "Hướng phát triển tương lai: Xây dựng cơ chế Machine Learning để tự học thói quen người dùng thay vì chỉ dựa vào Rule-based truyền thống.")

    # TÀI LIỆU THAM KHẢO
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_para(doc, "[1]. Tài liệu Spring Framework Boot, Security, AspectJ.\n"
                  "[2]. Android Developers: WorkManager, Navigation Component, Room Database.\n"
                  "[3]. Kiến trúc và Mô hình bảo mật Data Loss Prevention toàn diện.")

    doc.save("BaoCao_A_Care_DLP_Detailed.docx")

if __name__ == "__main__":
    create_report()
