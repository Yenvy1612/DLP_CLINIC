import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_code(doc, text):
    p = doc.add_paragraph(text, style='Macro Text')
    return p

def create_report():
    doc = Document()

    # Title Page
    title = doc.add_heading("BÁO CÁO DỰ ÁN TỐT NGHIỆP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("HỆ THỐNG QUẢN LÝ PHÒNG KHÁM A-CARE TÍCH HỢP GIẢI PHÁP PHÒNG CHỐNG THẤT THOÁT DỮ LIỆU (DLP)", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5):
        doc.add_paragraph()
    doc.add_page_break()

    # CHƯƠNG 1: MỞ ĐẦU
    add_heading(doc, "CHƯƠNG 1: MỞ ĐẦU", 1)
    
    add_heading(doc, "1.1. Bối cảnh đề tài", 2)
    add_para(doc, "Trong kỷ nguyên số hóa hiện nay, ngành y tế đang trải qua những thay đổi mang tính bước ngoặt. Việc chuyển đổi từ hồ sơ bệnh án giấy sang hồ sơ bệnh án điện tử (EMR - Electronic Medical Records) mang lại vô vàn lợi ích: dễ dàng lưu trữ, truy xuất thông tin nhanh chóng, hỗ trợ chuẩn đoán từ xa và tối ưu hóa quy trình quản lý bệnh viện/phòng khám. Tuy nhiên, sự tiện lợi này luôn song hành cùng những rủi ro an ninh mạng nghiêm trọng.")
    add_para(doc, "Dữ liệu y tế bao gồm các thông tin cực kỳ nhạy cảm: thông tin định danh cá nhân (PII - Personally Identifiable Information) như Họ tên, Căn cước công dân (CCCD), Số điện thoại, Email; và thông tin sức khỏe cá nhân (PHI - Protected Health Information) như nhóm máu, tiền sử bệnh lý, kết quả chẩn đoán các bệnh lý nhạy cảm (như HIV, Ung thư). Nếu những dữ liệu này bị rò rỉ, hậu quả không chỉ là tổn thất tài chính mà còn ảnh hưởng nặng nề đến danh dự, quyền riêng tư của bệnh nhân, đồng thời khiến cơ sở y tế đối mặt với các án phạt pháp lý nặng nề.")
    add_para(doc, "Nguyên nhân rò rỉ dữ liệu không chỉ đến từ các cuộc tấn công mạng bên ngoài (Hacker, Ransomware) mà một phần rất lớn đến từ các rủi ro nội bộ (Insider Threats). Nhân viên y tế, bác sĩ có thể vô tình hoặc cố ý sao chép, tải xuống, chụp ảnh màn hình, hoặc xuất file PDF bệnh án để chia sẻ ra ngoài hệ thống quản lý an toàn. Các biện pháp bảo mật truyền thống như tường lửa (Firewall) hay mã hóa dữ liệu tĩnh là chưa đủ để kiểm soát hành vi của người dùng hợp lệ bên trong hệ thống.")
    add_para(doc, "Trước bối cảnh đó, Giải pháp Phòng chống Thất thoát Dữ liệu (DLP - Data Loss Prevention) nổi lên như một mảnh ghép bảo mật thiết yếu. DLP giúp theo dõi, phát hiện và ngăn chặn các hành vi vi phạm chính sách bảo mật dữ liệu theo thời gian thực.")

    add_heading(doc, "1.2. Lý do chọn đề tài", 2)
    add_para(doc, "Việc xây dựng một hệ thống phòng khám thông thường hiện nay đã khá phổ biến. Tuy nhiên, một hệ thống tích hợp sâu cơ chế bảo mật DLP ở cả tầng máy chủ (Backend) lẫn tầng thiết bị đầu cuối (Endpoint/Mobile) là một bài toán khó và mang tính thực tiễn cao.")
    add_para(doc, "Đề tài 'Hệ thống Quản lý phòng khám A-Care tích hợp giải pháp phòng chống thất thoát dữ liệu (DLP)' được lựa chọn vì các lý do sau:")
    add_bullet(doc, "Giải quyết bài toán thực tế: Đáp ứng nhu cầu bảo vệ quyền riêng tư của bệnh nhân trong các phòng khám vừa và nhỏ, nơi mà hệ thống bảo mật thường bị xem nhẹ.")
    add_bullet(doc, "Tuân thủ tiêu chuẩn: Hướng đến việc tuân thủ các quy định quốc tế về bảo vệ dữ liệu y tế như HIPAA (Health Insurance Portability and Accountability Act) hay GDPR (General Data Protection Regulation).")
    add_bullet(doc, "Kiểm soát rủi ro nội bộ: Đặc biệt tập trung vào việc giám sát và hạn chế quyền hạn của Bác sĩ và Nhân viên - những người có quyền truy cập hợp lệ nhưng có nguy cơ làm rò rỉ dữ liệu qua các thao tác Copy/Paste, Export PDF trên thiết bị di động.")
    add_bullet(doc, "Tính mới của công nghệ: Triển khai Mobile Agent chạy ngầm trên Android để kiểm duyệt dữ liệu tại chỗ (Client-side DLP) kết hợp cùng hệ thống phân tích hành vi (User Behavior Analytics) sử dụng AOP (Aspect-Oriented Programming) trên Spring Boot Backend.")

    add_heading(doc, "1.3. Mục tiêu của đề tài", 2)
    add_para(doc, "Đề tài hướng tới việc hoàn thành các mục tiêu cụ thể sau đây:")
    add_bullet(doc, "Về mặt nghiệp vụ: Xây dựng nền tảng Mobile App và Backend phục vụ các chức năng cơ bản của một phòng khám: Quản lý người dùng (Bệnh nhân, Bác sĩ), Đặt lịch khám, Quản lý hồ sơ bệnh án, Xem lịch sử và xuất báo cáo.")
    add_bullet(doc, "Về mặt bảo mật Endpoint (Mobile Agent): Phát triển một Module Agent tích hợp trực tiếp vào Android App. Agent này có khả năng tự động đăng ký thiết bị, đồng bộ chính sách bảo mật (Policy Syncing) từ server, và chặn/che giấu (Masking) các thông tin nhạy cảm (CCCD, SĐT, Email, Keyword bệnh lý) khi người dùng thực hiện thao tác Copy hoặc Export PDF.")
    add_bullet(doc, "Về mặt bảo mật Backend (Network DLP & UBA): Xây dựng cơ chế đánh giá hành vi HTTP (HTTP Behavior Check) thông qua các Annotation (AOP). Hệ thống Backend có thể chặn các API truy xuất dữ liệu nếu phát hiện người dùng thao tác ngoài giờ hành chính, vượt quá số lượng request cho phép (Rate Limit) hoặc tải quá lượng dữ liệu trong ngày (Volume Limit).")
    add_bullet(doc, "Về mặt quản lý rủi ro (Audit & Logs): Thiết lập hệ thống thu thập sự kiện bất thường (DlpLogs và SecurityEvents), kết hợp với Room Database và WorkManager trên Mobile để đảm bảo log không bị mất ngay cả khi thiết bị mất kết nối mạng. Admin có thể xem toàn bộ trên Dashboard.")

    add_heading(doc, "1.4. Đối tượng và phạm vi nghiên cứu", 2)
    add_para(doc, "Đối tượng nghiên cứu:")
    add_bullet(doc, "Quy trình vận hành, đặt lịch và quản lý hồ sơ tại phòng khám đa khoa.")
    add_bullet(doc, "Các kỹ thuật nhận dạng mẫu (Regex Matching), xử lý chuỗi và che lấp dữ liệu (Data Masking).")
    add_bullet(doc, "Công nghệ phân tích luồng thực thi thông qua Aspect-Oriented Programming (AOP).")
    add_bullet(doc, "Quản lý tiến trình ngầm (Background Tasks) và đồng bộ hóa dữ liệu ngoại tuyến (Offline Syncing) trên nền tảng Android.")

    add_para(doc, "Phạm vi nghiên cứu:")
    add_bullet(doc, "Hệ thống chỉ tập trung vào nền tảng Android (đối với Client) và Java Spring Boot (đối với Server). Không bao gồm ứng dụng iOS.")
    add_bullet(doc, "Cơ chế DLP tập trung vào kiểm soát luồng dữ liệu dưới dạng văn bản (Text). Không đi sâu vào xử lý nhận diện hình ảnh (OCR) hay mã hóa toàn bộ ổ cứng thiết bị di động (MDM - Mobile Device Management).")
    add_bullet(doc, "Các luật (Rules) được thiết lập tĩnh (Rule-based) trên Backend, có thể cấu hình thông qua Admin. Chưa áp dụng các thuật toán Deep Learning phức tạp để tự động sinh luật.")

    add_heading(doc, "1.5. Phương pháp nghiên cứu", 2)
    add_para(doc, "Để hoàn thành đồ án, các phương pháp nghiên cứu sau được áp dụng:")
    add_bullet(doc, "Phương pháp nghiên cứu lý thuyết: Tìm hiểu tài liệu, các bài báo khoa học về DLP, HIPAA, UBA, kiến trúc Microservices, và các best-practices trong lập trình Android, Spring Boot.")
    add_bullet(doc, "Phương pháp mô hình hóa: Thiết kế sơ đồ Use Case, sơ đồ tuần tự (Sequence Diagram), sơ đồ luồng dữ liệu (Data Flow Diagram) và thiết kế cơ sở dữ liệu (ERD) trước khi tiến hành viết mã.")
    add_bullet(doc, "Phương pháp thực nghiệm (Thực hành): Trực tiếp cài đặt (Coding) ứng dụng dựa trên các tài liệu đã thiết kế. Sử dụng các framework hiện đại như Kotlin Coroutines, Retrofit, Room Database cho Mobile; và Hibernate, Spring Security, JWT cho Backend.")
    add_bullet(doc, "Phương pháp kiểm thử (Testing): Đặt ra các kịch bản (Scenarios) kiểm thử giả lập hành vi vi phạm (ví dụ: dùng tài khoản bác sĩ copy dữ liệu CCCD bệnh nhân vào ban đêm) để đánh giá độ trễ và khả năng phản ứng của hệ thống DLP.")

    add_heading(doc, "1.6. Cấu trúc của báo cáo", 2)
    add_para(doc, "Báo cáo được chia làm 5 chương:")
    add_bullet(doc, "Chương 1: Mở đầu - Trình bày tổng quan về bối cảnh, lý do, mục tiêu và phạm vi của đề tài.")
    add_bullet(doc, "Chương 2: Cơ sở lý thuyết và công nghệ sử dụng - Giới thiệu các khái niệm bảo mật cốt lõi và các công cụ, ngôn ngữ được sử dụng để xây dựng hệ thống.")
    add_bullet(doc, "Chương 3: Phân tích và Thiết kế hệ thống - Đi sâu vào chi tiết kiến trúc, các luồng nghiệp vụ (Flows), sơ đồ cơ sở dữ liệu, và thiết kế của Mobile Agent cũng như BehaviorAspect trên Backend.")
    add_bullet(doc, "Chương 4: Cài đặt và thử nghiệm - Mô tả quá trình triển khai mã nguồn, kết quả của các kịch bản kiểm thử (Test Cases) thực tế trên ứng dụng.")
    add_bullet(doc, "Chương 5: Tổng kết và đánh giá - Rút ra kết luận về những kết quả đạt được, điểm mạnh, điểm yếu và định hướng phát triển mở rộng trong tương lai.")

    doc.add_page_break()

    # CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG
    add_heading(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    
    add_heading(doc, "2.1. Tổng quan về Data Loss Prevention (DLP)", 2)
    add_para(doc, "DLP (Data Loss Prevention) là một chiến lược và giải pháp phần mềm nhằm đảm bảo dữ liệu nhạy cảm không bị truy cập trái phép, sử dụng sai mục đích hoặc rò rỉ ra bên ngoài phạm vi của tổ chức. Hệ thống DLP thường hoạt động dựa trên các quy tắc (policies) được xác định trước, liên tục phân tích và giám sát dòng chảy của dữ liệu.")
    add_para(doc, "Trong lĩnh vực y tế, DLP đặc biệt quan trọng do tính chất nhạy cảm của hồ sơ bệnh án (EMR). Các giải pháp DLP thường được chia làm ba loại chính:")
    add_bullet(doc, "Endpoint DLP: Cài đặt trực tiếp trên thiết bị của người dùng (máy tính, điện thoại). Nó kiểm soát các hành động cục bộ như copy/paste văn bản, in ấn, chụp ảnh màn hình, lưu file ra USB hoặc xuất PDF. (Trong đề tài này, Mobile Agent đóng vai trò là Endpoint DLP).")
    add_bullet(doc, "Network DLP: Được triển khai tại các gateway của mạng lưới. Nó giám sát lưu lượng mạng (Email, Web traffic, API calls) để phát hiện các gói tin chứa dữ liệu nhạy cảm gửi ra bên ngoài. (Trong đề tài này, Spring Boot Filter và AspectJ đóng vai trò tương tự Network DLP nội bộ).")
    add_bullet(doc, "Storage DLP / Cloud DLP: Quét và mã hóa các tệp tin lưu trữ tĩnh trên máy chủ hoặc đám mây (Data at Rest) để tìm kiếm các tệp bị chia sẻ sai quyền hạn.")

    add_heading(doc, "2.2. Phân tích Hành vi Người dùng (UBA - User Behavior Analytics)", 2)
    add_para(doc, "DLP truyền thống dựa trên nhận diện mẫu văn bản (Regex) có điểm yếu là chỉ phản ứng khi dữ liệu đã được yêu cầu xuất ra. Để nâng cao tính chủ động, UBA (Phân tích hành vi người dùng) được áp dụng. UBA thu thập lịch sử truy cập và đánh giá rủi ro dựa trên ngữ cảnh (Context-aware).")
    add_para(doc, "Hệ thống A-Care áp dụng UBA thông qua cơ chế Rule-Engine tại Backend:")
    add_bullet(doc, "Time-based Context: Phân tích múi giờ truy cập. Bác sĩ truy cập hồ sơ bệnh án vào lúc 2 giờ sáng là một dấu hiệu bất thường nghiêm trọng (Anomaly).")
    add_bullet(doc, "Volume-based Context: Phân tích lưu lượng. Một bác sĩ tải 100 hồ sơ bệnh án trong 5 phút thay vì 1-2 hồ sơ như bình thường là dấu hiệu của hành vi trộm cắp dữ liệu hàng loạt (Data Exfiltration).")
    add_bullet(doc, "Velocity / Rate-limit: Số lần gọi API liên tục từ một thiết bị trong thời gian cực ngắn báo hiệu thiết bị đó có thể đã bị chiếm quyền bởi Bot/Script.")

    add_heading(doc, "2.3. Data Masking và Mã hóa", 2)
    add_para(doc, "Data Masking (Che giấu dữ liệu) là kỹ thuật thay thế một phần dữ liệu nhạy cảm bằng các ký tự vô nghĩa (thường là dấu `*` hoặc `X`) nhằm che giấu thông tin thật nhưng vẫn giữ nguyên định dạng của văn bản để không làm thay đổi luồng xử lý hoặc cấu trúc hiển thị của ứng dụng.")
    add_para(doc, "Ví dụ: Số Căn cước công dân '034123456789' sẽ được mask thành '034******789'. Việc Masking được thực hiện ngay tại Client (Mobile App) đảm bảo rằng dữ liệu nhạy cảm hoàn toàn không xuất hiện trên giao diện xuất PDF của bác sĩ, ngay cả khi API từ server vẫn trả về dữ liệu thô (để phục vụ bệnh nhân xem hồ sơ của chính mình).")

    add_heading(doc, "2.4. Nền tảng phát triển Android (Kotlin & Architecture Components)", 2)
    add_para(doc, "Ứng dụng Client được xây dựng trên hệ điều hành Android, sử dụng ngôn ngữ Kotlin. Kotlin mang lại sự an toàn (Null-safety) và cú pháp ngắn gọn, kết hợp hoàn hảo với Kotlin Coroutines để xử lý đa luồng (Asynchronous programming).")
    add_para(doc, "Hệ thống tuân thủ kiến trúc MVVM (Model-View-ViewModel) giúp tách biệt hoàn toàn giao diện (UI) và logic dữ liệu (Business Logic). Các thư viện cốt lõi được sử dụng bao gồm:")
    add_bullet(doc, "Retrofit 2 & OkHttp: Quản lý giao tiếp mạng, gọi API lên server Spring Boot. Interceptor được cấu hình để tự động gắn JWT Token vào header của mọi request.")
    add_bullet(doc, "Room Database: Giải pháp ORM (Object-Relational Mapping) mạnh mẽ trên Android, dùng làm kho lưu trữ cục bộ (Local Storage). Room đặc biệt quan trọng đối với EventQueueRepository của Mobile Agent, giúp lưu lại các Log vi phạm khi thiết bị mất kết nối internet.")
    add_bullet(doc, "WorkManager: API chính thức của Google để xử lý các tác vụ nền tảng (Background processing) một cách đáng tin cậy. WorkManager được dùng cho `HeartbeatWorker` (báo cáo trạng thái thiết bị mỗi 15 phút) và `EventSyncWorker` (đồng bộ log vi phạm lên server khi có mạng).")
    add_bullet(doc, "Navigation Component: Quản lý luồng chuyển màn hình (Routing) bên trong Single Activity (MainActivity), kết hợp với BottomNavigationView tạo ra trải nghiệm mượt mà.")

    add_heading(doc, "2.5. Nền tảng Backend (Java Spring Boot)", 2)
    add_para(doc, "Backend là trái tim của hệ thống, xử lý toàn bộ logic nghiệp vụ (Auth, Appointment, Medical Record) và nhận luồng Log từ các thiết bị Android. Spring Boot được lựa chọn nhờ hệ sinh thái đồ sộ và độ ổn định cao trong ứng dụng doanh nghiệp.")
    add_bullet(doc, "Spring Security & JWT: Đảm nhiệm xác thực (Authentication) và phân quyền (Authorization) dựa trên Token không trạng thái (Stateless). Người dùng được chia thành 3 vai trò rõ rệt: PATIENT, DOCTOR, ADMIN.")
    add_bullet(doc, "Spring Data JPA & Hibernate: Ánh xạ thực thể (Entity) trong mã Java xuống các bảng trong cơ sở dữ liệu MySQL.")
    add_bullet(doc, "Aspect-Oriented Programming (Spring AOP): Áp dụng mẫu thiết kế hướng khía cạnh. Thay vì viết mã kiểm tra quyền hạn và ghi log thủ công ở mọi hàm Controller, AOP cho phép định nghĩa các `Aspect` (như `BehaviorAspect`). Các hàm API cần bảo vệ chỉ cần gắn annotation `@DlpProtected`. Khi có request, `BehaviorAspect` sẽ xen ngang (Intercept) luồng thực thi, kiểm tra các bộ luật của DLP (Giờ giấc, Volume) trước khi cho phép hàm API chạy tiếp.")
    
    doc.add_page_break()

    # CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
    add_heading(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    
    add_heading(doc, "3.1. Tổng quan kiến trúc hệ thống", 2)
    add_para(doc, "Hệ thống A-Care DLP Clinic hoạt động dựa trên mô hình Client-Server phân tán, nhưng có thêm một lớp Security trung gian tại Client. Kiến trúc bao gồm 3 khối chính:")
    add_para(doc, "1. Clinic Mobile App (User Interface): Là giao diện tương tác với Bệnh nhân và Bác sĩ. Đảm nhiệm các chức năng đăng ký, đặt lịch, xem hồ sơ.")
    add_para(doc, "2. Mobile Agent (Client-side DLP): Một Module chạy ngầm bên trong cấu trúc của Mobile App. Nó đóng vai trò như một người gác cổng trên thiết bị, liên tục lấy các chính sách (Policy) từ Server, quét (Scan) dữ liệu nhạy cảm trên bộ nhớ tạm (Clipboard) hoặc nội dung màn hình (Forms), chặn hành vi trích xuất dữ liệu, và đẩy sự kiện (Event) về Server.")
    add_para(doc, "3. Spring Boot Backend (Network DLP & Core Logic): Lưu trữ Database, cung cấp API nghiệp vụ, và thực hiện phân tích hành vi HTTP (HTTP Behavior Check). Nó cũng tiếp nhận toàn bộ các Log từ hàng ngàn thiết bị Mobile gộp lại, lưu vào `dlp_logs` để Admin kiểm soát trên Web Dashboard.")

    add_heading(doc, "3.2. Cấu trúc cơ sở dữ liệu (Database Schema)", 2)
    add_para(doc, "Cơ sở dữ liệu MySQL được thiết kế bao gồm các bảng nghiệp vụ và các bảng phục vụ bảo mật.")
    
    add_heading(doc, "3.2.1. Các bảng Nghiệp vụ", 3)
    add_bullet(doc, "Bảng `users`: Lưu trữ thông tin đăng nhập, mật khẩu (mã hóa BCrypt), Role (PATIENT, DOCTOR, ADMIN), họ tên, email, CCCD (patient_id_number).")
    add_bullet(doc, "Bảng `appointments`: Quản lý các lịch khám. Liên kết Foreign Key tới `patient_id` và `doctor_id`. Các trường gồm start_time, status (PENDING, DONE, CANCELLED), reason.")
    add_bullet(doc, "Bảng `medical_records`: Hồ sơ bệnh án. Được sinh ra sau khi bác sĩ khám xong lịch hẹn. Gồm các thông tin như diagnosis (chẩn đoán), clinical_notes (ghi chú), treatment_plan (phác đồ). Đây là nguồn dữ liệu nhạy cảm chính cần được bảo vệ.")

    add_heading(doc, "3.2.2. Các bảng Bảo mật và Agent", 3)
    add_bullet(doc, "Bảng `agents`: Lưu thông tin của Mobile Agent. Bao gồm `device_id` (MAC hoặc Android ID), `platform` (ANDROID), `os_version`, `status` (ONLINE/OFFLINE), `last_heartbeat`.")
    add_bullet(doc, "Bảng `agent_events`: (Nằm trên SQLite của Room Mobile Database, và cũng được đồng bộ lên MySQL) Gồm `event_type` (VIEW, COPY, EXPORT, MATCHED), `severity` (LOW, HIGH, CRITICAL), `content_snippet` (văn bản đã được Masked chứa vi phạm).")
    add_bullet(doc, "Bảng `dlp_logs`: Lưu vết toàn bộ các cảnh báo vi phạm. Được tổng hợp từ Agent Events và cả từ Backend (AOP vi phạm rules). Cung cấp lịch sử bằng chứng để truy tố nếu cần thiết.")

    add_heading(doc, "3.3. Các Luồng nghiệp vụ cốt lõi (Core Business Flows)", 2)
    
    add_heading(doc, "3.3.1. Luồng Khởi động và Splash Screen", 3)
    add_para(doc, "Ngay khi icon Ứng dụng được người dùng nhấn vào, hệ điều hành Android tạo tiến trình và khởi tạo lớp Application. Tại `App.kt`, hệ thống gọi `AgentInitializer.init()` để đưa Agent vào trạng thái sẵn sàng. Module này khởi tạo `AgentDatabase` (Room), `RetrofitClient` cho Agent API (chạy song song với API nghiệp vụ), `DlpScanner` và `PolicyManager`.")
    add_para(doc, "Sau đó, `SplashActivity` bắt đầu. Nó truy cập `SessionManager` để kiểm tra có token lưu trữ không. Nếu token còn hiệu lực, app bỏ qua màn hình đăng nhập và chuyển thẳng vào `MainActivity` (Bottom Navigation). Ngược lại, luồng chuyển hướng về `LoginActivity`.")

    add_heading(doc, "3.3.2. Luồng Đăng nhập (Login) và Đăng ký Thiết bị", 3)
    add_para(doc, "Tại `LoginActivity`, người dùng nhập thông tin. Nút Đăng nhập được bấm sẽ kích hoạt luồng gọi API `POST /api/auth/login`. Nếu backend trả về Status 200 OK cùng chuỗi JWT, hệ thống thực hiện 4 bước cực kỳ quan trọng đồng thời:")
    add_para(doc, "1. Lưu JWT vào EncryptedSharedPreferences (SessionManager).")
    add_para(doc, "2. Gọi `AgentManager.register(username)`: Thiết bị Android tự sinh một chuỗi ID duy nhất (Android_ID) gửi lên endpoint `POST /api/agents/register` để báo với Server rằng 'Thiết bị này đang gắn với tài khoản Bác sĩ A'.")
    add_para(doc, "3. Gọi `AgentManager.syncPolicy()`: Tải các mẫu Regex, danh sách từ khóa cấm (HIV, Ung thư, Tuyệt mật) từ `GET /api/agents/policy`. Danh sách này được nạp vào bộ nhớ (Cache) của `PolicyManager` để dùng cho quét cục bộ.")
    add_para(doc, "4. Gọi `AgentManager.sendHeartbeat()`: Báo cáo thiết bị đang ONLINE. Kích hoạt WorkManager chạy lặp lại mỗi 15 phút.")

    add_heading(doc, "3.3.3. Luồng Đặt lịch khám (Appointment Booking)", 3)
    add_para(doc, "Dành cho PATIENT. Người dùng truy cập tab Lịch hẹn, chọn Bác sĩ, chuyên khoa và khung thời gian. Khi bấm xác nhận, `BookingActivity` gom thông tin tạo payload `AppointmentCreateRequest` đẩy lên `POST /api/appointments`. Tại tầng Backend, `AppointmentService` kiểm tra xung đột thời gian (Conflict), nếu khung giờ đó chưa có ai đặt, lịch sẽ được tạo. Đồng thời, Mobile Agent gửi nhẹ một event Audit `eventType=CREATE_APPOINTMENT`, `severity=LOW` về hệ thống log để ghi nhận hành vi.")

    add_heading(doc, "3.4. Đi sâu vào cơ chế Mobile Agent (Endpoint DLP)", 2)
    
    add_heading(doc, "3.4.1. Luồng DLP Scan trong Form và Text", 3)
    add_para(doc, "Bất cứ khi nào Bác sĩ nhập dữ liệu vào ô mô tả, hoặc thao tác Copy một văn bản từ App, ứng dụng sẽ chuyển chuỗi văn bản (Text) đó qua `DlpScanner.scan(text)`. Hàm `scan()` duyệt qua toàn bộ các mẫu Regex (CCCD 9-12 số, Email, SĐT) và Keywords có trong Policy.")
    add_para(doc, "Nếu kết quả trả về `MATCHED`, ứng dụng ngay lập tức chặn hành động Copy (Clipboard không được cấp quyền). Sau đó, chuỗi văn bản chứa vi phạm sẽ được chạy qua `MaskingUtil.mask()` để thay các con số thành dấu `***` (VD: 'Bệnh nhân có CCCD 034123456789' trở thành 'Bệnh nhân có CCCD 034******789'). Chuỗi văn bản đã an toàn này được lưu vào đối tượng `AgentEvent` kèm `violationType=CCCD` và đẩy xuống `EventQueueRepository`.")

    add_heading(doc, "3.4.2. Luồng xử lý Export PDF an toàn", 3)
    add_para(doc, "Đây là luồng phức tạp nhất nhằm cân bằng giữa tiện ích và bảo mật. Khi bác sĩ nhấn nút 'Xuất danh sách Lịch khám và Hồ sơ'.")
    add_para(doc, "1. Mobile App tập hợp dữ liệu. Đối với mỗi bệnh án, nó lấy các trường dữ liệu nhạy cảm (patientIdNumber, email, phone).")
    add_para(doc, "2. Thay vì dùng regex quét rủi ro trên toàn văn bản PDF gây tốn hiệu năng, ứng dụng thực hiện Hard-Masking (Ép che giấu cố định): Nó trích xuất chính xác các biến `patientIdNumber`, `email`, và `phone` và can thiệp bằng logic thay chuỗi trực tiếp thành các ký tự `***`. Mọi CCCD dù có format sai lệ hay ngắn cũng đều biến mất hoàn toàn trên góc nhìn của tệp xuất ra.")
    add_para(doc, "3. Sau khi dữ liệu được Mask, toàn bộ nội dung text được gửi một lệnh Track tới `AgentEventTracker.trackExport()`. Kèm theo event `EXPORT_PDF`, `severity=LOW` hoặc `HIGH` tùy cấu hình.")
    add_para(doc, "4. Cuối cùng, thư viện `PdfExportUtil` vẽ chuỗi văn bản đã an toàn lên các trang PDF (Canvas) và lưu vào thư mục Downloads của điện thoại.")

    add_heading(doc, "3.4.3. Luồng Queue Offline (Room Database & WorkManager)", 3)
    add_para(doc, "Điểm yếu của đa số hệ thống giám sát là khi điện thoại mất mạng (Offline), các log bị mất, dẫn đến việc người xấu cố tình tắt 4G/Wifi trước khi copy dữ liệu.")
    add_para(doc, "Hệ thống A-Care giải quyết bằng mô hình Local-Queue. Khi `AgentEventTracker` bắt được lỗi (VD Block Copy), nó không gọi API ngay, mà gọi `AgentEventDao.insert()` lưu bản ghi vào SQLite thông qua Room. Bản ghi có cờ `is_synced = false`.")
    add_para(doc, "Thư viện Android WorkManager chứa lớp `EventSyncWorker`. Worker này được hệ điều hành đánh thức khi và chỉ khi có mạng (Constraints: CONNECTED). Khi chạy, nó truy vấn `SELECT * FROM agent_events WHERE is_synced = false`. Gom thành danh sách (Batch) và gọi `POST /api/agent-events`. Nhận được Status 200 từ Backend, nó cập nhật cờ `is_synced = true` hoặc xóa log cục bộ. Nhờ vậy, không một sự kiện vi phạm nào bị bỏ lọt.")

    add_heading(doc, "3.5. Đi sâu vào cơ chế Network DLP và Phân tích hành vi (Backend)", 2)
    
    add_heading(doc, "3.5.1. Aspect-Oriented Programming (AOP) và `BehaviorAspect`", 3)
    add_para(doc, "Trong thư mục `backend/src/main/java/com/acare/backend/common`, có một lớp đặc biệt mang tên `BehaviorAspect`. Thay vì phải chèn các đoạn code if/else ở khắp các file Controller để chặn người dùng truy cập ngoài giờ, hệ thống định nghĩa annotation `@DlpProtected(action=\"VIEW\")`.")
    add_para(doc, "Tại hàm `getMedicalRecordByPatientId()`, ta chỉ cần gắn annotation trên. Spring AOP sẽ tự động bọc (wrap) hàm này bằng phương thức `checkDlpRules(ProceedingJoinPoint joinPoint, DlpProtected dlpProtected)` bên trong `BehaviorAspect`. Lúc này, hệ thống sẽ kiểm tra ngữ cảnh (Context):")
    add_para(doc, "- Bước 1: Trích xuất `UserId` và `Role` từ chuỗi JWT.")
    add_para(doc, "- Bước 2: Kiểm tra `isWorkingHour()`. Giả sử hệ thống định nghĩa giờ hành chính là 07:00 - 22:00. Nếu request đến lúc 02:00 sáng. AOP sẽ xác minh Role. Nếu là Bệnh nhân (PATIENT), hệ thống cho phép qua (vì bệnh nhân có quyền xem bệnh án của chính mình 24/7). Nếu là Bác sĩ (DOCTOR), hàm lập tức `throw new SecurityException(\"DLP: Truy cập bị từ chối. Ngoài giờ làm việc.\")`. Đồng thời, một bản ghi `DlpLog` với `RiskLevel.HIGH` được sinh ra lưu vào Database.")
    
    add_heading(doc, "3.5.2. Chống Bot và Cào Dữ liệu (Rate Limiting)", 3)
    add_para(doc, "Cũng tại `BehaviorAspect`, hệ thống gọi hàm `ruleEngineService.isRateLimitExceeded(userId)`. Bằng cách tận dụng Redis, mỗi khi có request lấy hồ sơ, key Redis `user_req:{userId}` sẽ tăng giá trị đếm (INCR). Nếu trong 1 phút, một tài khoản Bác sĩ cố tình kéo tới 50 hồ sơ bệnh án (vượt ngưỡng), hệ thống chặn đứng với mã lỗi 429 (Too Many Requests), ghi log `RATE_LIMIT` mức độ `CRITICAL` và có thể tự động khóa tài khoản.")

    doc.add_page_break()

    # CHƯƠNG 4: CÀI ĐẶT VÀ THỬ NGHIỆM
    add_heading(doc, "CHƯƠNG 4: CÀI ĐẶT VÀ THỬ NGHIỆM", 1)
    
    add_heading(doc, "4.1. Môi trường triển khai", 2)
    add_para(doc, "Hệ thống được phát triển trên môi trường Windows, IDE Android Studio Ladybug cho Mobile và IntelliJ IDEA cho Backend. Mobile App sử dụng Gradle build tool phiên bản 8.x. Backend chạy trên JDK 17, Spring Boot 3.2, kết nối MySQL 8.0.")
    add_para(doc, "Quá trình biên dịch Android được thực thi bằng lệnh `gradlew installDebug`, tự động cài đặt tệp APK lên thiết bị di động thật hoặc máy ảo Emulator.")

    add_heading(doc, "4.2. Thử nghiệm các Kịch bản nghiệp vụ", 2)
    
    add_heading(doc, "4.2.1. Kịch bản 1: Giám sát quá trình xuất dữ liệu y tế (Export PDF) của Bác sĩ", 3)
    add_para(doc, "Mục đích: Đảm bảo dữ liệu CCCD, SĐT của bệnh nhân không bị lọt ra file PDF khi Bác sĩ tải về thiết bị.")
    add_para(doc, "Tiến hành: Bác sĩ đăng nhập thành công vào ứng dụng. Di chuyển đến danh sách Lịch khám. Nhấn nút 'Export PDF'. Ứng dụng tiến hành gom danh sách hồ sơ.")
    add_para(doc, "Kết quả: File PDF được sinh ra trong thư mục Downloads của điện thoại. Khi mở file, phần thông tin bệnh nhân hiển thị dưới dạng: `CCCD=***, Email=tes***@gmail.com, SĐT=037***59`. Đồng thời, một log `EXPORT_PDF` được ghi nhận trên Backend. Tính riêng tư của bệnh nhân được bảo vệ tuyệt đối.")

    add_heading(doc, "4.2.2. Kịch bản 2: Truy cập dữ liệu ngoài giờ (Time-based Anomaly)", 3)
    add_para(doc, "Mục đích: Đảm bảo nhân viên y tế không thể lén lút lấy hồ sơ vào đêm khuya.")
    add_para(doc, "Tiến hành: Đổi giờ hệ thống máy chủ sang 01:30 AM. Trên thiết bị di động, Bác sĩ bấm vào tab Hồ sơ để tải danh sách bệnh án.")
    add_para(doc, "Kết quả: Ứng dụng lập tức nhận về kết quả lỗi HTTP 403 / 500 do `SecurityException` từ `BehaviorAspect`. Màn hình hiển thị thông báo lỗi, danh sách hồ sơ trống rỗng. Tại bảng điều khiển quản trị (Admin Dashboard), lập tức xuất hiện cảnh báo ĐỎ (CRITICAL) chỉ định rõ Bác sĩ nào, IP nào đã cố gắng vi phạm.")
    add_para(doc, "Thử nghiệm phụ: Cùng lúc đó, sử dụng tài khoản Bệnh nhân (PATIENT) để truy cập tab Hồ sơ của chính mình. Kết quả: Bệnh nhân vẫn tải và xem hồ sơ bình thường. Bộ lọc Role trong AOP đã hoạt động chuẩn xác.")

    add_heading(doc, "4.2.3. Kịch bản 3: Xử lý Log ngoại tuyến (Offline Event Queueing)", 3)
    add_para(doc, "Mục đích: Đánh giá khả năng bắt vi phạm khi không có Internet.")
    add_para(doc, "Tiến hành: Tắt hoàn toàn Wifi và 3G trên điện thoại. Bác sĩ mở một hồ sơ lưu trong Cache (hoặc thực hiện nhập form). Bác sĩ nhập từ khóa cấm 'HIV' và bấm gửi (Submit).")
    add_para(doc, "Kết quả: DlpScanner chặn thao tác, tạo cảnh báo vi phạm. Do không có mạng, event được chèn vào bảng SQLite `agent_events` với trạng thái `is_synced = 0`. Hai phút sau, bật Wifi lại. `EventSyncWorker` của Android WorkManager được đánh thức, ngay lập tức quét DB, lấy event và đẩy lên `POST /api/agent-events`. Backend nhận log, trạng thái cục bộ chuyển thành `is_synced = 1`. Kiểm thử hoàn toàn thành công.")

    doc.add_page_break()

    # CHƯƠNG 5: TỔNG KẾT VÀ ĐÁNH GIÁ
    add_heading(doc, "CHƯƠNG 5: TỔNG KẾT VÀ ĐÁNH GIÁ", 1)
    
    add_heading(doc, "5.1. Kết quả đạt được", 2)
    add_para(doc, "Đề tài đã hoàn thành việc xây dựng và tích hợp thành công hai khía cạnh khó nhất của bảo mật dữ liệu: Endpoint DLP (Mobile Agent) và Network DLP (Backend AOP). Hệ thống A-Care không chỉ chứng minh khả năng quản lý tốt các nghiệp vụ y tế hằng ngày (đặt lịch, lưu bệnh án) mà còn sở hữu vòng bảo mật kép vững chắc.")
    add_para(doc, "Mobile Agent trên Android đã chứng minh khả năng hoạt động ngầm ổn định, bắt chính xác các sự kiện, che lấp (Masking) dữ liệu nhạy cảm cực kỳ hiệu quả mà không làm giảm tốc độ trải nghiệm của người dùng. Backend Aspect-Oriented Programming đã xử lý mượt mà việc phân tích ngữ cảnh (thời gian, role, lưu lượng) trước khi trả dữ liệu về mạng.")

    add_heading(doc, "5.2. Hạn chế còn tồn tại", 2)
    add_para(doc, "Mặc dù đạt được nhiều thành quả, hệ thống vẫn còn một số điểm cần khắc phục:")
    add_bullet(doc, "Thuật toán quét Regex trong DlpScanner đôi khi vẫn bị phụ thuộc vào định dạng chuẩn. Nếu người dùng nhập CCCD chèn khoảng trắng (VD: 034 123 456), regex truyền thống có thể bỏ lọt (dù đã áp dụng cơ chế Hard-Masking riêng cho hàm Export để khắc phục).")
    add_bullet(doc, "Backend vẫn dựa vào các ngưỡng tĩnh (Static Thresholds) để chặn Rate Limit hoặc Time-based. Điều này khiến cấu hình hệ thống cứng nhắc.")

    add_heading(doc, "5.3. Hướng phát triển trong tương lai", 2)
    add_para(doc, "Trong các phiên bản tiếp theo, dự án có thể mở rộng theo các hướng sau:")
    add_bullet(doc, "Tích hợp Trí tuệ nhân tạo (AI/Machine Learning) vào Module Anomaly Detection. AI sẽ tự động học (Learn) thói quen đăng nhập, truy cập dữ liệu của từng bác sĩ cụ thể. Nếu một bác sĩ thường làm việc đêm, AI sẽ không đánh dấu đó là rủi ro cao, nhưng nếu bác sĩ tải quá số lượng, AI sẽ điều chỉnh điểm Risk Level tự động (Adaptive Rules).")
    add_bullet(doc, "Bổ sung cơ chế Watermark (Đóng dấu chìm) động trên màn hình Mobile. Ngay cả khi bác sĩ dùng máy ảnh ngoài để chụp lén màn hình điện thoại, các dấu chấm nhỏ li ti chứa thông tin ID Bác sĩ (Invisible Watermarking) sẽ giúp truy vết thủ phạm rò rỉ.")
    add_bullet(doc, "Mở rộng hệ sinh thái Agent sang nền tảng iOS (Swift) và Web Browser Extension.")

    doc.add_page_break()

    # TÀI LIỆU THAM KHẢO
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_para(doc, "[1]. Craig Walls (2020), Spring in Action, 6th Edition, Manning Publications.")
    add_para(doc, "[2]. Android Developers Documentation (2024), WorkManager and Room Persistence Library, Google.")
    add_para(doc, "[3]. The Health Insurance Portability and Accountability Act (HIPAA), U.S. Department of Health & Human Services.")
    add_para(doc, "[4]. Gartner Research (2023), Magic Quadrant for Enterprise Data Loss Prevention.")
    add_para(doc, "[5]. Baeldung (2024), Intro to Spring AOP (Aspect-Oriented Programming).")

    doc.save("BaoCao_A_Care_DLP_50_Pages.docx")
    print("Massive detailed report generated.")

if __name__ == "__main__":
    create_report()
